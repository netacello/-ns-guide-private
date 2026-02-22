"""
Build the Neurosteer Physician Guide as a formatted .docx file.
Reads the markdown source and figures, outputs a professional Word document.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# from docx.oxml.ns import qn  # not needed
from docx.shared import Cm, Inches, Pt, RGBColor

BASEDIR = Path(__file__).parent
MD_PATH = BASEDIR / "Neurosteer_Physician_Guide_v6.md"
FIG_DIR = BASEDIR / "physician_guide_figures"
OUT_PATH = BASEDIR / "Neurosteer_Physician_Guide_v6.docx"

# Colors
NEUROSTEER_BLUE = RGBColor(0x1A, 0x52, 0x76)
DARK_GRAY = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM_GRAY = RGBColor(0x7F, 0x8C, 0x8D)


def setup_styles(doc):
    """Configure document styles for a professional medical document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = DARK_GRAY
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15

    # Heading 1 — Main sections
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = NEUROSTEER_BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — Subsections
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = NEUROSTEER_BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3 — Sub-subsections
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = NEUROSTEER_BLUE
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True


def add_formatted_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    return run


def parse_inline_formatting(paragraph, text):
    """Parse markdown bold/italic and add as formatted runs."""
    # Split on bold+italic patterns
    parts = re.split(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            add_formatted_run(paragraph, part[3:-3], bold=True, italic=True)
        elif part.startswith("**") and part.endswith("**"):
            add_formatted_run(paragraph, part[2:-2], bold=True)
        elif part.startswith("*") and part.endswith("*"):
            add_formatted_run(paragraph, part[1:-1], italic=True)
        else:
            paragraph.add_run(part)


def add_table_from_md(doc, lines):
    """Parse a markdown table and add it to the document."""
    # Filter out separator lines (|---|---|)
    data_lines = [l for l in lines if not re.match(r"^\|[\s\-:|]+\|$", l)]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 1:
        return

    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = NEUROSTEER_BLUE

    # Set column widths to be more compact
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.0)

    doc.add_paragraph()  # spacing after table


def process_markdown(doc, md_text):
    """Process markdown text and build the Word document."""
    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            if in_table and table_lines:
                add_table_from_md(doc, table_lines)
                table_lines = []
                in_table = False
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_" * 80)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(6)
            i += 1
            continue

        # Title (# level 1)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            # Main document title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = NEUROSTEER_BLUE
            i += 1
            continue

        # Heading 2 (##)
        if stripped.startswith("## ") and not stripped.startswith("### "):
            text = stripped[3:].strip()
            if text == "Table of Contents":
                # Skip TOC — Word can generate its own
                i += 1
                while (
                    i < len(lines)
                    and lines[i].strip()
                    and not lines[i].strip().startswith("## ")
                ):
                    i += 1
                continue
            # Check if subtitle (like "Comprehensive Physician Guide")
            if "Physician Guide" in text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(text)
                run.font.size = Pt(14)
                run.font.color.rgb = MEDIUM_GRAY
                run.italic = True
            else:
                doc.add_heading(text, level=1)
            i += 1
            continue

        # Heading 3 (###)
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            text = stripped[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # Heading 4 (####)
        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Image
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            # Resolve path
            full_path = BASEDIR / img_path
            if full_path.exists():
                doc.add_picture(str(full_path), width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                add_formatted_run(
                    p, f"[Figure: {alt_text}]", italic=True, color=MEDIUM_GRAY
                )
            i += 1
            continue

        # Figure caption (starts with *)
        if stripped.startswith("*") and stripped.endswith("*") and "Figure" in stripped:
            text = stripped.strip("*").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = MEDIUM_GRAY
            i += 1
            continue

        # Table line
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue

        # Flush any pending table
        if in_table and table_lines:
            add_table_from_md(doc, table_lines)
            table_lines = []
            in_table = False

        # Bold metadata lines (Version, For use by)
        if stripped.startswith("**Version") or stripped.startswith("**For use by"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = stripped.replace("**", "")
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = MEDIUM_GRAY
            i += 1
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r"^(\d+)\.\s+", stripped)
        if num_match:
            text = stripped[num_match.end() :].strip()
            p = doc.add_paragraph(style="List Number")
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Note/caveat lines starting with *Note:
        if stripped.startswith("*Note:") or stripped.startswith("*Note "):
            text = stripped.strip("*").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            run.font.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = MEDIUM_GRAY
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)
        i += 1

    # Flush any remaining table
    if table_lines:
        add_table_from_md(doc, table_lines)


def main():
    print("Building physician guide .docx ...")

    # Read markdown
    md_text = MD_PATH.read_text(encoding="utf-8")

    # Create document
    doc = Document()

    # Page setup — Letter, with reasonable margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)

    # Setup styles
    setup_styles(doc)

    # Add header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Neurosteer Brain Metrics Assessment \u2014 Physician Guide v6.0")
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.italic = True

    # Simple footer (no XML field injection to avoid corruption)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Neurosteer Ltd. \u2014 Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_GRAY

    # Process the markdown content
    process_markdown(doc, md_text)

    # Save
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
